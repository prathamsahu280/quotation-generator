"""renderer.py — render a list of "blocks" onto a letterhead .docx and optionally to PDF.

A spec is a dict {"blocks": [...]}. Block types: paragraph, heading, bullet,
spacer, hr, table. Headers, footers and margins of the letterhead are preserved;
blocks are appended after any existing body content.

This is a refactor of the original plugin script `scripts/generate_quotation.py`
into an importable module used by the web backend.
"""
from __future__ import annotations

import os
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "centre": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _set_cell_shading(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color="000000", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:color"), color)
        tc_borders.append(e)
    tc_pr.append(tc_borders)


def _apply_run_style(run, style: dict):
    if not style:
        return
    if "font" in style:
        run.font.name = style["font"]
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), style["font"])
    if "size" in style:
        run.font.size = Pt(style["size"])
    if "bold" in style:
        run.bold = bool(style["bold"])
    if "italic" in style:
        run.italic = bool(style["italic"])
    if "underline" in style:
        run.underline = bool(style["underline"])
    if "color" in style:
        c = style["color"].lstrip("#")
        run.font.color.rgb = RGBColor.from_string(c.upper())


def _apply_para_style(para, style: dict):
    if not style:
        return
    if "alignment" in style:
        align = ALIGN_MAP.get(str(style["alignment"]).lower())
        if align is not None:
            para.alignment = align
    pf = para.paragraph_format
    if "space_before" in style:
        pf.space_before = Pt(style["space_before"])
    if "space_after" in style:
        pf.space_after = Pt(style["space_after"])
    if "line_spacing" in style:
        pf.line_spacing = float(style["line_spacing"])
    if "indent_left" in style:
        pf.left_indent = Inches(style["indent_left"])
    if "keep_with_next" in style:
        pf.keep_with_next = bool(style["keep_with_next"])


def _add_text_with_breaks(para, text: str, style: dict):
    parts = str(text).split("\n")
    for i, part in enumerate(parts):
        if i > 0:
            para.add_run().add_break()
        run = para.add_run(part)
        _apply_run_style(run, style)


def add_paragraph_block(doc, block):
    style = block.get("style", {})
    para = doc.add_paragraph()
    _apply_para_style(para, style)
    _add_text_with_breaks(para, block.get("text", ""), style)
    return para


def add_heading_block(doc, block):
    style = block.get("style", {})
    level = int(block.get("level", 1))
    para = doc.add_paragraph()
    _apply_para_style(para, style)
    base_size = style.get("size", {1: 16, 2: 14, 3: 12}.get(level, 12))
    eff_style = {"bold": style.get("bold", True), "size": base_size, **style}
    _add_text_with_breaks(para, block.get("text", ""), eff_style)
    return para


def add_bullet_block(doc, block):
    style = block.get("style", {})
    level = int(block.get("level", 0))
    bullet_style = "List Bullet" if level == 0 else f"List Bullet {min(level + 1, 3)}"
    try:
        para = doc.add_paragraph(style=bullet_style)
    except KeyError:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
        block["text"] = "• " + block.get("text", "")
    _apply_para_style(para, style)
    _add_text_with_breaks(para, block.get("text", ""), style)
    return para


def add_spacer_block(doc, block):
    for _ in range(int(block.get("lines", 1))):
        doc.add_paragraph()


def add_hr_block(doc, block):
    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(block.get("size", 6)))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), block.get("color", "auto"))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_table_block(doc, block):
    style = block.get("style", {})
    headers = block.get("headers", [])
    rows = block.get("rows", [])
    n_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
    if n_cols == 0:
        return None

    n_rows = (1 if headers else 0) + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = False
    if style.get("table_style"):
        try:
            table.style = style["table_style"]
        except KeyError:
            pass

    col_widths = block.get("column_widths_in")
    if col_widths and len(col_widths) == n_cols:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)

    cell_font = style.get("font")
    cell_size = style.get("size", 10)
    borders = style.get("borders", True)
    header_align = style.get("header_align", "center")
    row_aligns = style.get("row_aligns") or ["left"] * n_cols

    if headers:
        for col_idx, h in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.text = ""
            _apply_para_style(para, {"alignment": header_align, "space_before": 2, "space_after": 2})
            rs = {"size": cell_size, "bold": style.get("header_bold", True)}
            if style.get("header_color"):
                rs["color"] = style["header_color"]
            if cell_font:
                rs["font"] = cell_font
            _add_text_with_breaks(para, str(h), rs)
            if borders:
                _set_cell_borders(cell)
            if style.get("header_shading"):
                _set_cell_shading(cell, style["header_shading"].lstrip("#").upper())

    start = 1 if headers else 0
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row[:n_cols]):
            cell = table.rows[start + r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.text = ""
            align = row_aligns[c_idx] if c_idx < len(row_aligns) else "left"
            _apply_para_style(para, {"alignment": align, "space_before": 2, "space_after": 2})
            rs = {"size": cell_size}
            if cell_font:
                rs["font"] = cell_font
            _add_text_with_breaks(para, str(val), rs)
            if borders:
                _set_cell_borders(cell)

    return table


BLOCK_RENDERERS = {
    "paragraph": add_paragraph_block,
    "heading": add_heading_block,
    "bullet": add_bullet_block,
    "spacer": add_spacer_block,
    "hr": add_hr_block,
    "table": add_table_block,
}


def render_blocks(doc, blocks):
    for block in blocks:
        renderer = BLOCK_RENDERERS.get(block.get("type"))
        if renderer is None:
            raise ValueError(f"Unknown block type: {block.get('type')!r}")
        renderer(doc, block)


def _find_soffice() -> str | None:
    """Locate a LibreOffice/OpenOffice headless binary (used on Linux/cloud)."""
    env = os.getenv("SOFFICE_BIN")
    if env and Path(env).exists():
        return env
    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    return None


def to_pdf(docx_path: Path) -> Path:
    """Convert a .docx to .pdf next to it.

    Prefers LibreOffice headless (works on a Linux server, no MS Word); falls
    back to docx2pdf (Windows/macOS + MS Word) for local use. Raises if neither
    is available so the caller can still hand back the .docx.
    """
    docx_path = Path(docx_path)
    pdf_path = docx_path.with_suffix(".pdf")

    soffice = _find_soffice()
    if soffice:
        # LibreOffice writes <stem>.pdf into --outdir; that matches pdf_path.
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir",
             str(docx_path.parent), str(docx_path)],
            check=True, capture_output=True, timeout=120,
        )
        if not pdf_path.exists():
            raise RuntimeError("LibreOffice did not produce a PDF.")
        return pdf_path

    # Fallback: docx2pdf (needs MS Word installed).
    from docx2pdf import convert as docx2pdf_convert
    docx2pdf_convert(str(docx_path), str(pdf_path))
    return pdf_path


def pdf_engine() -> str:
    """Human-readable name of the active PDF backend (for diagnostics)."""
    if _find_soffice():
        return "libreoffice"
    try:
        import docx2pdf  # noqa: F401
        return "docx2pdf"
    except ImportError:
        return "none"


def render(letterhead_path: Path, spec: dict, out_docx: Path, make_pdf: bool = True):
    """Render `spec` onto a copy of the letterhead. Returns (docx_path, pdf_path|None)."""
    out_docx = Path(out_docx)
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(letterhead_path, out_docx)

    doc = Document(str(out_docx))
    render_blocks(doc, spec.get("blocks", []))
    doc.save(str(out_docx))

    pdf_path = to_pdf(out_docx) if make_pdf else None
    return out_docx, pdf_path
