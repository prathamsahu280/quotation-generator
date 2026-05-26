"""pdf_letterhead.py — turn a PDF letterhead into a .docx we can render onto.

A PDF letterhead (often a scanned/designed full-page image) can't be appended to
like a normal Word letterhead: a direct PDF→.docx conversion dumps the whole page
into the document *body*, so any quotation we add spills onto a blank second page.

Instead we rasterise the first PDF page and place it as a **full-page background
image in the document header** (behind text, repeating on every page), leaving
the body empty. The quotation then renders on top of the letterhead. We also
auto-detect the header/footer ink bands so the body margins keep text inside the
blank middle area.

The page is rasterised with **LibreOffice headless** (`soffice --convert-to png`),
the same binary renderer.py already uses for .docx→PDF — so a Linux/KVM deploy
needs no extra native libraries beyond LibreOffice. Page dimensions come from
`pypdf`.

build_spec/renderer are unchanged: they just append body blocks, which now flow
over the background.
"""
from __future__ import annotations

import subprocess
import tempfile
from copy import deepcopy
from pathlib import Path

import numpy as np
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Emu
from PIL import Image
from pypdf import PdfReader

import renderer  # reuse _find_soffice() so PDF and DOCX paths share one binary

PT_TO_EMU = 12700          # 1 pt = 1/72 in; 1 in = 914400 EMU
IN_TO_EMU = 914400


def _render_pdf_page1_png(pdf_path: Path, out_png: Path) -> Path:
    """Rasterise the first PDF page to PNG using LibreOffice headless.

    LibreOffice writes <stem>.png into --outdir; we move it to out_png. Raises
    if LibreOffice is unavailable or produces nothing, so the caller can surface
    a clear "try a .docx instead" message.
    """
    soffice = renderer._find_soffice()
    if not soffice:
        raise RuntimeError(
            "LibreOffice (soffice) not found; install libreoffice or set SOFFICE_BIN."
        )
    outdir = out_png.parent
    subprocess.run(
        [soffice, "--headless", "--convert-to", "png", "--outdir",
         str(outdir), str(pdf_path)],
        check=True, capture_output=True, timeout=120,
    )
    produced = outdir / (pdf_path.stem + ".png")
    if not produced.exists():
        raise RuntimeError("LibreOffice did not produce a PNG from the PDF.")
    if produced != out_png:
        produced.replace(out_png)
    return out_png


def _detect_margins(gray: np.ndarray, dpi: float) -> tuple[float, float, float, float]:
    """Return (top, bottom, left, right) margins in inches that clear the
    letterhead's printed header/footer/side bands, from a grayscale page."""
    H, W = gray.shape
    ink = gray < 200                       # dark-ish pixels = printed content
    row_ink = ink.sum(axis=1)
    row_thr = max(3, int(W * 0.005))
    rows = np.where(row_ink > row_thr)[0]

    def clamp(v, lo, hi):
        return max(lo, min(hi, v))

    # header: last content row in the top 45% of the page
    top_rows = rows[rows < 0.45 * H] if rows.size else np.array([])
    top_in = (top_rows.max() / dpi + 0.18) if top_rows.size else 0.8
    top = clamp(top_in, 0.6, 3.0)

    # footer: first content row in the bottom 40%
    bot_rows = rows[rows > 0.60 * H] if rows.size else np.array([])
    bot_in = ((H - bot_rows.min()) / dpi + 0.18) if bot_rows.size else 0.8
    bottom = clamp(bot_in, 0.5, 2.5)

    # sides: content columns within the middle band only
    mid = ink[int(0.25 * H): int(0.75 * H), :]
    col_thr = max(3, int(mid.shape[0] * 0.01))
    cols = np.where(mid.sum(axis=0) > col_thr)[0]
    left = clamp((cols.min() / dpi + 0.1) if cols.size else 0.9, 0.5, 1.3)
    right = clamp(((W - cols.max()) / dpi + 0.1) if cols.size else 0.9, 0.5, 1.3)
    return top, bottom, left, right


def _make_background_anchor(inline_drawing) -> None:
    """Convert the <wp:inline> image (just added to the header) into a page-
    anchored, behind-text <wp:anchor> so it becomes a full-page background."""
    inline = inline_drawing.find(qn("wp:inline"))
    extent = inline.find(qn("wp:extent"))
    graphic = inline.find(qn("a:graphic"))
    cx, cy = extent.get("cx"), extent.get("cy")

    anchor = parse_xml(
        f'<wp:anchor {nsdecls("wp", "a", "r", "pic")} behindDoc="1" distT="0" '
        'distB="0" distL="0" distR="0" simplePos="0" locked="0" layoutInCell="1" '
        'allowOverlap="1" relativeHeight="0">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="100" name="Letterhead background"/>'
        "</wp:anchor>"
    )
    anchor.append(deepcopy(graphic))
    inline.getparent().replace(inline, anchor)


def pdf_to_letterhead_docx(pdf_path: Path, out_docx: Path) -> Path:
    """Render PDF page 1 as a full-page background in a new .docx's header."""
    pdf_path, out_docx = Path(pdf_path), Path(out_docx)

    # page 1 dimensions in points (pypdf); account for any /Rotate
    page = PdfReader(str(pdf_path)).pages[0]
    w_pt, h_pt = float(page.mediabox.width), float(page.mediabox.height)
    if (page.rotation or 0) % 180 == 90:
        w_pt, h_pt = h_pt, w_pt
    w_emu = int(round(w_pt * PT_TO_EMU))
    h_emu = int(round(h_pt * PT_TO_EMU))

    with tempfile.TemporaryDirectory(prefix="lh_") as tmp:
        png = _render_pdf_page1_png(pdf_path, Path(tmp) / "page.png")

        arr = np.asarray(Image.open(png).convert("RGB"))
        gray = arr.mean(axis=2)
        # derive the effective DPI from the rendered pixel width vs page width
        dpi = arr.shape[1] / (w_pt / 72.0)
        top, bottom, left, right = _detect_margins(gray, dpi)

        doc = Document()
        sec = doc.sections[0]
        sec.page_width = Emu(w_emu)
        sec.page_height = Emu(h_emu)
        sec.top_margin = Emu(int(top * IN_TO_EMU))
        sec.bottom_margin = Emu(int(bottom * IN_TO_EMU))
        sec.left_margin = Emu(int(left * IN_TO_EMU))
        sec.right_margin = Emu(int(right * IN_TO_EMU))
        sec.header_distance = Emu(0)
        sec.footer_distance = Emu(0)

        # full-page image in the header, then make it a behind-text page anchor
        hp = sec.header.paragraphs[0]
        run = hp.add_run()
        run.add_picture(str(png), width=Emu(w_emu), height=Emu(h_emu))
        drawing = run._element.find(qn("w:drawing"))
        _make_background_anchor(drawing)

        doc.save(str(out_docx))
    return out_docx
